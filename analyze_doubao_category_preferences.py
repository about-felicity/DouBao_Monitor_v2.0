import csv
import json
import math
import os
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path

import doubao_dashboard_server as dashboard


BASE_DIR = Path(__file__).resolve().parent
REPORT_DIR = BASE_DIR / "reports"
REPORT_DIR.mkdir(exist_ok=True)


CATEGORY_RULES = (
    ("染发剂", ("染发",)),
    ("眉毛精华液", ("眉毛",)),
    ("睫毛精华液", ("睫毛",)),
    ("美白面霜", ("美白面霜",)),
    ("控油蓬松洗发水", ("控油蓬松洗发水",)),
    ("二硫化硒洗发水", ("二硫化硒",)),
    ("防断洗发水", ("防断洗发水",)),
    ("防脱洗发水", ("防脱洗发水",)),
    ("防脱精华液", ("防脱精华液",)),
    ("护发精油", ("护发精油",)),
    ("护发素", ("护发素",)),
    ("防晒霜", ("防晒霜",)),
    ("沐浴精油", ("沐浴精油",)),
    ("面膜", ("面膜",)),
    ("祛痘精华液", ("祛痘",)),
    ("护手霜", ("护手霜",)),
    ("洗面奶", ("洗面奶",)),
    ("身体乳", ("身体乳",)),
    ("造型喷雾", ("造型喷雾",)),
    ("爽肤水", ("爽肤水",)),
    ("眼霜", ("眼霜",)),
    ("卸妆油", ("卸妆油",)),
)


def category_for(question):
    text = str(question or "").strip()
    for category, tokens in CATEGORY_RULES:
        if any(token in text for token in tokens):
            return category
    return text or "未分类"


def intent_for(question):
    text = str(question or "")
    return "评价类" if "评价" in text or "怎么样" in text else "推荐类"


def pct(value, total, digits=1):
    return round(value * 100 / total, digits) if total else 0.0


def fmt_pct(value):
    return f"{value:.1f}%"


def clean_cell(value):
    return str(value or "").replace("|", "／").replace("\n", " ").strip()


def md_table(headers, rows):
    lines = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in rows:
        lines.append("| " + " | ".join(clean_cell(item) for item in row) + " |")
    return "\n".join(lines)


def rank_position(counter, name):
    ordered = sorted(counter.items(), key=lambda item: (-item[1], item[0]))
    previous = None
    rank = 0
    for position, (candidate, count) in enumerate(ordered, 1):
        if count != previous:
            rank = position
            previous = count
        if candidate == name:
            return rank
    return None


def daily_record(rows, seen_before):
    refs = len(rows)
    runs = {str(row.get("run_no") or "").strip() for row in rows if str(row.get("run_no") or "").strip()}
    urls = {str(row.get("href") or "").strip() for row in rows if str(row.get("href") or "").strip()}
    type_counter = Counter(row["_source_type"] for row in rows)
    media_counter = Counter(row["_media"] for row in rows)
    url_counter = Counter(str(row.get("href") or "").strip() for row in rows if str(row.get("href") or "").strip())
    top_media, top_media_count = media_counter.most_common(1)[0] if media_counter else ("-", 0)
    return {
        "refs": refs,
        "runs": len(runs),
        "urls": urls,
        "unique_urls": len(urls),
        "type_counter": type_counter,
        "media_counter": media_counter,
        "video_share": pct(type_counter.get("视频", 0), refs),
        "article_share": pct(type_counter.get("文章", 0), refs),
        "product_share": pct(type_counter.get("商品页", 0), refs),
        "top_media": top_media,
        "top_media_share": pct(top_media_count, refs),
        "top10_concentration": pct(sum(count for _url, count in url_counter.most_common(10)), refs),
        "new_url_rate": pct(len(urls - seen_before), len(urls)) if seen_before else None,
        "questions": sorted({row["_question"] for row in rows}),
    }


def transition_event(category, prev_day, day, prev, current, latest_day):
    union = prev["urls"] | current["urls"]
    intersection = prev["urls"] & current["urls"]
    jaccard = len(intersection) / len(union) if union else 1.0
    all_types = set(prev["type_counter"]) | set(current["type_counter"])
    mix_shift = 0.5 * sum(
        abs(pct(prev["type_counter"].get(t, 0), prev["refs"], 3) - pct(current["type_counter"].get(t, 0), current["refs"], 3))
        for t in all_types
    )
    turnover = (1 - jaccard) * 100
    new_rate = current["new_url_rate"] or 0.0
    concentration_shift = abs(current["top10_concentration"] - prev["top10_concentration"])
    score = round(0.35 * mix_shift + 0.35 * turnover + 0.20 * new_rate + 0.10 * concentration_shift, 1)
    question_mix_changed = set(prev["questions"]) != set(current["questions"])
    top_media_changed = prev["top_media"] != current["top_media"]
    severe = (
        score >= 45
        or mix_shift >= 25
        or jaccard <= 0.35
        or new_rate >= 55
        or (top_media_changed and mix_shift >= 15)
    )
    level = "剧烈" if severe else ("明显" if score >= 25 else "平稳")
    reasons = []
    if mix_shift >= 15:
        reasons.append(f"信源类型结构移动 {mix_shift:.1f}pct")
    if jaccard <= 0.55:
        reasons.append(f"链接集合相似度仅 {jaccard:.2f}")
    if new_rate >= 35:
        reasons.append(f"新链接率 {new_rate:.1f}%")
    if top_media_changed:
        reasons.append(f"主导媒体由 {prev['top_media']} 变为 {current['top_media']}")
    if question_mix_changed:
        reasons.append("该品类当天运行的问题组合发生变化")
    if day == latest_day:
        reasons.append("最新日仍可能继续增长")
    return {
        "category": category,
        "from_day": prev_day,
        "to_day": day,
        "level": level,
        "score": score,
        "jaccard": round(jaccard, 3),
        "turnover": round(turnover, 1),
        "mix_shift": round(mix_shift, 1),
        "new_url_rate": round(new_rate, 1),
        "concentration_shift": round(concentration_shift, 1),
        "video_shift": round(current["video_share"] - prev["video_share"], 1),
        "article_shift": round(current["article_share"] - prev["article_share"], 1),
        "top_media_from": prev["top_media"],
        "top_media_to": current["top_media"],
        "question_mix_changed": question_mix_changed,
        "reason": "；".join(reasons) or "主要指标变化较小",
    }


def preference_text(item):
    types = item["type_shares"]
    video = types.get("视频", 0)
    article = types.get("文章", 0)
    product = types.get("商品页", 0)
    if video >= 85:
        medium = "极强视频偏好"
    elif video >= 70:
        medium = "视频主导"
    elif video >= 55:
        medium = "视频优先、文章补充"
    elif article >= 35:
        medium = "视频与文章并重"
    else:
        medium = "混合型信源结构"
    if item["top10_share"] >= 55:
        pool = "高度依赖少数固定链接"
    elif item["top10_share"] >= 30:
        pool = "存在明显核心链接池"
    else:
        pool = "信源较分散、轮换较多"
    extra = []
    if product >= 5:
        extra.append(f"商品页占比偏高（{product:.1f}%）")
    if item["evaluation_refs"] and item["recommendation_refs"]:
        diff = item["intent_video_shares"].get("推荐类", 0) - item["intent_video_shares"].get("评价类", 0)
        extra.append(f"推荐类视频占比较评价类高 {diff:.1f}pct")
    tail = "；" + "；".join(extra) if extra else ""
    return f"{medium}，{pool}，主导媒体为 {item['top_media'][0][0] if item['top_media'] else '-'}{tail}。"


def analyze():
    rows = [row for row in dashboard.read_csv_rows() if not dashboard.is_quarantined_source_row(row)]
    ai_cache = dashboard.read_json(dashboard.AI_CACHE_PATH)
    meta_cache = dashboard.read_json(dashboard.META_CACHE_PATH)
    for row in rows:
        question = dashboard.question_for(row)
        source_type, media, host, _note = dashboard.source_for(row, ai_cache, meta_cache)
        row["_question"] = question
        row["_category"] = category_for(question)
        row["_intent"] = intent_for(question)
        row["_day"] = dashboard.date_for(row)
        row["_source_type"] = source_type
        row["_media"] = media
        row["_host"] = host

    snapshot = datetime.now(dashboard.CST).strftime("%Y-%m-%d %H:%M:%S")
    latest_day = max((row["_day"] for row in rows if row["_day"]), default="")
    by_category = defaultdict(list)
    for row in rows:
        by_category[row["_category"]].append(row)

    categories = []
    all_events = []
    for category, category_rows in by_category.items():
        refs = len(category_rows)
        runs = {str(row.get("run_no") or "").strip() for row in category_rows if str(row.get("run_no") or "").strip()}
        urls = {str(row.get("href") or "").strip() for row in category_rows if str(row.get("href") or "").strip()}
        type_counter = Counter(row["_source_type"] for row in category_rows)
        media_counter = Counter(row["_media"] for row in category_rows)
        url_counter = Counter(str(row.get("href") or "").strip() for row in category_rows if str(row.get("href") or "").strip())
        url_runs = defaultdict(set)
        title_by_url = {}
        for row in category_rows:
            href = str(row.get("href") or "").strip()
            if href:
                url_runs[href].add(str(row.get("run_no") or "").strip())
                title_by_url.setdefault(href, str(row.get("title") or "").strip())

        by_intent = defaultdict(list)
        for row in category_rows:
            by_intent[row["_intent"]].append(row)
        intent_video_shares = {
            intent: pct(sum(1 for row in intent_rows if row["_source_type"] == "视频"), len(intent_rows))
            for intent, intent_rows in by_intent.items()
        }

        day_rows = defaultdict(list)
        for row in category_rows:
            if row["_day"]:
                day_rows[row["_day"]].append(row)
        daily = {}
        seen_before = set()
        for day in sorted(day_rows):
            daily[day] = daily_record(day_rows[day], seen_before)
            seen_before |= daily[day]["urls"]
        events = []
        days = sorted(daily)
        for prev_day, day in zip(days, days[1:]):
            event = transition_event(category, prev_day, day, daily[prev_day], daily[day], latest_day)
            events.append(event)
            all_events.append(event)

        top_links = []
        for href, count in url_counter.most_common(5):
            top_links.append({
                "title": title_by_url.get(href) or href,
                "href": href,
                "count": count,
                "run_count": len(url_runs[href]),
                "run_coverage": pct(len(url_runs[href]), len(runs)),
            })
        top10_share = pct(sum(count for _url, count in url_counter.most_common(10)), refs)
        type_shares = {name: pct(count, refs) for name, count in type_counter.items()}
        item = {
            "category": category,
            "questions": sorted({row["_question"] for row in category_rows}),
            "refs": refs,
            "runs": len(runs),
            "unique_urls": len(urls),
            "refs_per_run": round(refs / len(runs), 2) if runs else 0,
            "type_counter": type_counter,
            "type_shares": type_shares,
            "top_media": media_counter.most_common(5),
            "top10_share": top10_share,
            "top_links": top_links,
            "daily": daily,
            "events": events,
            "severe_events": [event for event in events if event["level"] == "剧烈"],
            "recommendation_refs": len(by_intent.get("推荐类", [])),
            "evaluation_refs": len(by_intent.get("评价类", [])),
            "intent_video_shares": intent_video_shares,
        }
        item["preference"] = preference_text(item)
        categories.append(item)

    categories.sort(key=lambda item: (-item["runs"], item["category"]))
    all_events.sort(key=lambda item: (-item["score"], item["category"], item["to_day"]))
    return rows, categories, all_events, snapshot, latest_day


def write_summary_csv(categories, path):
    fields = [
        "品类", "包含问题", "运行轮次", "信源条数", "唯一链接", "平均每轮信源", "视频占比", "文章占比",
        "商品页占比", "第一媒体", "第一媒体次数", "前10链接集中度", "剧烈变化次数", "豆包偏好结论",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for item in categories:
            top_media = item["top_media"][0] if item["top_media"] else ("-", 0)
            writer.writerow({
                "品类": item["category"],
                "包含问题": "；".join(item["questions"]),
                "运行轮次": item["runs"],
                "信源条数": item["refs"],
                "唯一链接": item["unique_urls"],
                "平均每轮信源": item["refs_per_run"],
                "视频占比": item["type_shares"].get("视频", 0),
                "文章占比": item["type_shares"].get("文章", 0),
                "商品页占比": item["type_shares"].get("商品页", 0),
                "第一媒体": top_media[0],
                "第一媒体次数": top_media[1],
                "前10链接集中度": item["top10_share"],
                "剧烈变化次数": len(item["severe_events"]),
                "豆包偏好结论": item["preference"],
            })


def write_events_csv(events, path):
    fields = ["品类", "前一观测日", "当前观测日", "等级", "波动分", "链接相似度", "信源结构移动pct", "新链接率", "视频变化pct", "文章变化pct", "原主导媒体", "新主导媒体", "原因"]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        for event in events:
            if event["level"] == "平稳":
                continue
            writer.writerow({
                "品类": event["category"], "前一观测日": event["from_day"], "当前观测日": event["to_day"],
                "等级": event["level"], "波动分": event["score"], "链接相似度": event["jaccard"],
                "信源结构移动pct": event["mix_shift"], "新链接率": event["new_url_rate"],
                "视频变化pct": event["video_shift"], "文章变化pct": event["article_shift"],
                "原主导媒体": event["top_media_from"], "新主导媒体": event["top_media_to"], "原因": event["reason"],
            })


def write_markdown(rows, categories, events, snapshot, latest_day, path):
    total_refs = len(rows)
    total_runs = len({str(row.get("run_no") or "").strip() for row in rows if str(row.get("run_no") or "").strip()})
    total_questions = len({row["_question"] for row in rows})
    total_urls = len({str(row.get("href") or "").strip() for row in rows if str(row.get("href") or "").strip()})
    overall_types = Counter(row["_source_type"] for row in rows)
    video_rank = sorted(categories, key=lambda item: item["type_shares"].get("视频", 0), reverse=True)
    article_rank = sorted(categories, key=lambda item: item["type_shares"].get("文章", 0), reverse=True)
    concentration_rank = sorted(categories, key=lambda item: item["top10_share"], reverse=True)
    severe = [event for event in events if event["level"] == "剧烈"]

    out = []
    out.append("# 豆包各产品品类信源偏好与变化波动报告")
    out.append("")
    out.append(f"**数据快照：** {snapshot}（中国标准时间）  ")
    out.append(f"**覆盖数据：** {total_runs:,} 轮、{total_refs:,} 条信源、{total_urls:,} 个唯一链接、{total_questions} 个标准化问题、{len(categories)} 个产品品类  ")
    out.append(f"**最新数据日：** {latest_day}（仍在运行，最新日指标属于未收盘数据）")
    out.append("")
    out.append("---")
    out.append("")
    out.append("## 一、结论摘要")
    out.append("")
    out.append(f"整体仍以视频为主：视频 {fmt_pct(pct(overall_types.get('视频', 0), total_refs))}、文章 {fmt_pct(pct(overall_types.get('文章', 0), total_refs))}、商品页 {fmt_pct(pct(overall_types.get('商品页', 0), total_refs))}。但不同品类差异明显，不能用全站平均替代单品类判断。")
    out.append("")
    out.append(f"视频偏好最强的品类包括：{'、'.join(item['category'] + '（' + fmt_pct(item['type_shares'].get('视频', 0)) + '）' for item in video_rank[:5])}。")
    out.append("")
    out.append(f"文章占比较高的品类包括：{'、'.join(item['category'] + '（' + fmt_pct(item['type_shares'].get('文章', 0)) + '）' for item in article_rank[:5])}。")
    out.append("")
    out.append(f"核心链接最集中的品类包括：{'、'.join(item['category'] + '（前10链接占' + fmt_pct(item['top10_share']) + '）' for item in concentration_rank[:5])}。集中度高意味着豆包长期依赖固定候选池，也意味着一两个核心链接变化会明显影响结果。")
    out.append("")
    out.append(f"共识别到 {len(severe)} 次“剧烈变化”。变化可能来自链接池真正切换，也可能来自当天运行的问题组合改变或最新日尚未跑完；报告逐项给出原因标签。")
    out.append("")
    out.append("## 二、统计口径与剧烈变化定义")
    out.append("")
    out.append("- 品类由同义问题合并，例如染发剂包含通用推荐、首迷评价、JSV 评价和梵玢评价；眉毛/睫毛精华液同时包含通用推荐与依斯佩尔评价问题。")
    out.append("- 运行轮次按不同 `run_no` 去重；一个运行轮次可包含多条信源。")
    out.append("- 视频/文章/商品页由现有 AI 信源缓存和规则共同判定。")
    out.append("- 链接相似度采用 Jaccard：相邻观测日共同链接数 ÷ 链接并集数。越低表示链接池变化越大。")
    out.append("- 信源结构移动为视频、文章、商品页等占比变化的总变差，单位为百分点。")
    out.append("- 波动分综合信源类型变化、链接换血、新链接率和核心链接集中度变化。满足波动分≥45、结构移动≥25pct、链接相似度≤0.35等条件时标为“剧烈”。")
    out.append("- 相邻观测日不一定是自然日连续；没有运行的日期不会制造空白比较。")
    out.append("")
    out.append("## 三、全部品类偏好总表")
    out.append("")
    summary_rows = []
    for item in categories:
        top_media = item["top_media"][0][0] if item["top_media"] else "-"
        summary_rows.append([
            item["category"], item["runs"], f"{item['refs']:,}", item["unique_urls"],
            fmt_pct(item["type_shares"].get("视频", 0)), fmt_pct(item["type_shares"].get("文章", 0)),
            fmt_pct(item["type_shares"].get("商品页", 0)), top_media, fmt_pct(item["top10_share"]),
            len(item["severe_events"]), item["preference"],
        ])
    out.append(md_table(["品类", "轮次", "信源", "唯一链接", "视频", "文章", "商品页", "主导媒体", "前10链接占比", "剧烈变化", "豆包偏好"], summary_rows))
    out.append("")
    out.append("## 四、变化最剧烈的时段")
    out.append("")
    if severe:
        event_rows = []
        for event in severe[:30]:
            event_rows.append([
                event["category"], f"{event['from_day']} → {event['to_day']}", event["score"],
                f"{event['jaccard']:.2f}", f"{event['mix_shift']:.1f}pct", fmt_pct(event["new_url_rate"]),
                f"{event['video_shift']:+.1f}pct", f"{event['article_shift']:+.1f}pct", event["reason"],
            ])
        out.append(md_table(["品类", "时段", "波动分", "链接相似度", "结构移动", "新链接率", "视频变化", "文章变化", "主要原因"], event_rows))
    else:
        out.append("当前没有达到剧烈变化阈值的时段。")
    out.append("")
    out.append("> 最新数据日尚未结束。凡变化终点为最新日的事件，都应在当日运行完成后再作最终判断。")
    out.append("")
    out.append("## 五、分品类详细分析")
    out.append("")

    for index, item in enumerate(categories, 1):
        out.append(f"### {index}. {item['category']}")
        out.append("")
        out.append(f"**包含问题：** {'、'.join(item['questions'])}  ")
        out.append(f"**豆包偏好：** {item['preference']}  ")
        out.append(f"**数据规模：** {item['runs']:,} 轮、{item['refs']:,} 条信源、{item['unique_urls']:,} 个唯一链接、平均每轮 {item['refs_per_run']:.2f} 条。")
        out.append("")
        out.append(md_table(
            ["视频", "文章", "商品页", "其他", "前10链接集中度", "剧烈变化次数"],
            [[fmt_pct(item["type_shares"].get("视频", 0)), fmt_pct(item["type_shares"].get("文章", 0)), fmt_pct(item["type_shares"].get("商品页", 0)), fmt_pct(item["type_shares"].get("其他", 0)), fmt_pct(item["top10_share"]), len(item["severe_events"])]],
        ))
        out.append("")
        if item["evaluation_refs"] and item["recommendation_refs"]:
            out.append(f"推荐类视频占比 {fmt_pct(item['intent_video_shares'].get('推荐类', 0))}，评价类视频占比 {fmt_pct(item['intent_video_shares'].get('评价类', 0))}。该差值反映问法变化，而不是品类本身突然改变。")
            out.append("")
        media_rows = [[name, f"{count:,}", fmt_pct(pct(count, item["refs"]))] for name, count in item["top_media"]]
        out.append("**主要媒体/平台**")
        out.append("")
        out.append(md_table(["媒体/平台", "引用次数", "品类内占比"], media_rows))
        out.append("")
        out.append("**核心链接**")
        out.append("")
        link_rows = [[link["title"][:70], link["count"], link["run_count"], fmt_pct(link["run_coverage"])] for link in item["top_links"]]
        out.append(md_table(["标题", "引用次数", "覆盖轮次", "运行覆盖率"], link_rows))
        out.append("")
        out.append("**每日信源变化**")
        out.append("")
        daily_rows = []
        previous = None
        for day, rec in item["daily"].items():
            jac = "-"
            if previous:
                union = previous["urls"] | rec["urls"]
                jac = f"{len(previous['urls'] & rec['urls']) / len(union):.2f}" if union else "1.00"
            daily_rows.append([
                day + ("（未收盘）" if day == latest_day else ""), rec["runs"], rec["refs"], rec["unique_urls"],
                fmt_pct(rec["video_share"]), fmt_pct(rec["article_share"]), fmt_pct(rec["product_share"]),
                rec["top_media"], fmt_pct(rec["new_url_rate"]) if rec["new_url_rate"] is not None else "-", jac,
                "、".join(rec["questions"]),
            ])
            previous = rec
        out.append(md_table(["日期", "轮次", "信源", "唯一链接", "视频", "文章", "商品页", "主导媒体", "新链接率", "与前日相似度", "当天问题"], daily_rows))
        out.append("")
        notable = [event for event in item["events"] if event["level"] != "平稳"]
        if notable:
            out.append("**明显/剧烈变化：**")
            out.append("")
            for event in sorted(notable, key=lambda event: -event["score"]):
                out.append(f"- {event['from_day']} → {event['to_day']}：**{event['level']}**（{event['score']} 分）。{event['reason']}。")
        elif len(item["daily"]) <= 1:
            out.append("**变化判断：** 目前只有一个观测日，无法判断日变化。")
        else:
            out.append("**变化判断：** 相邻观测日没有达到明显波动阈值。")
        out.append("")

    out.append("## 六、业务使用建议")
    out.append("")
    out.append("1. 每个品类单独设基线，不要用全站视频占比评价某个品类。")
    out.append("2. 品牌趋势使用“品牌出现轮次 ÷ 当日运行轮次”，不能使用品牌提及总数直接比较。")
    out.append("3. 剧烈变化必须拆解为：问题组合变化、媒体结构变化、核心链接消失、新链接进入、最新日未收盘。")
    out.append("4. 对链接集中度高的品类重点监控前 10 个核心链接；对链接分散的品类重点监控新链接率和主导媒体变化。")
    out.append("5. 推荐类与评价类分开展示。评价问题通常文章和商品页更多，合并后会制造虚假的品类波动。")
    out.append("6. 最新日完成后重新生成报告，只有连续两个完整观测日仍保持变化，才标记为稳定趋势。")
    out.append("")
    out.append("> 本报告分析的是豆包在当前美妆个护监控问题集中的信源选择偏好，不代表豆包在所有领域的全局偏好。")
    path.write_text("\n".join(out), encoding="utf-8")


def markdown_to_docx(md_path, docx_path):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt, RGBColor
    except Exception:
        return False
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(9.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for name, size, color in (("Title", 22, "173F35"), ("Heading 1", 16, "0B806A"), ("Heading 2", 13, "146B5B"), ("Heading 3", 11, "2F665D")):
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    def clean(text):
        text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
        text = re.sub(r"`([^`]*)`", r"\1", text)
        return text.strip()

    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if not line or line == "---":
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip()):
            headers = [clean(value) for value in line.strip("|").split("|")]
            index += 2
            data = []
            while index < len(lines) and lines[index].startswith("|"):
                data.append([clean(value) for value in lines[index].strip("|").split("|")])
                index += 1
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Light Shading Accent 1"
            for col, value in enumerate(headers):
                table.rows[0].cells[col].text = value
            for row in data:
                cells = table.add_row().cells
                for col in range(len(headers)):
                    cells[col].text = row[col] if col < len(row) else ""
            for table_row in table.rows:
                for cell in table_row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Microsoft YaHei"
                            run.font.size = Pt(7.5)
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            continue
        heading = re.match(r"^(#{1,3})\s+(.*)$", line)
        if heading:
            level = len(heading.group(1))
            text = clean(heading.group(2))
            if level == 1 and not any(paragraph.text for paragraph in doc.paragraphs):
                paragraph = doc.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(text)
            else:
                doc.add_heading(text, level=level)
            index += 1
            continue
        if line.startswith("> "):
            doc.add_paragraph(clean(line[2:]), style="Intense Quote")
            index += 1
            continue
        if re.match(r"^[-*]\s+", line):
            doc.add_paragraph(clean(re.sub(r"^[-*]\s+", "", line)), style="List Bullet")
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(clean(re.sub(r"^\d+\.\s+", "", line)), style="List Number")
            index += 1
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.add_run(clean(line))
        index += 1
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("豆包各产品品类信源偏好与变化波动报告")
    doc.save(docx_path)
    return True


def main():
    rows, categories, events, snapshot, latest_day = analyze()
    stamp = latest_day or datetime.now(dashboard.CST).strftime("%Y-%m-%d")
    md_path = REPORT_DIR / f"豆包各产品品类偏好与波动报告_{stamp}.md"
    docx_path = md_path.with_suffix(".docx")
    summary_path = REPORT_DIR / f"豆包各产品品类偏好汇总_{stamp}.csv"
    events_path = REPORT_DIR / f"豆包品类剧烈变化明细_{stamp}.csv"
    write_markdown(rows, categories, events, snapshot, latest_day, md_path)
    write_summary_csv(categories, summary_path)
    write_events_csv(events, events_path)
    markdown_to_docx(md_path, docx_path)
    print(json.dumps({
        "snapshot": snapshot,
        "latest_day": latest_day,
        "categories": len(categories),
        "events": len(events),
        "severe_events": sum(1 for event in events if event["level"] == "剧烈"),
        "markdown": str(md_path),
        "docx": str(docx_path),
        "summary_csv": str(summary_path),
        "events_csv": str(events_path),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
